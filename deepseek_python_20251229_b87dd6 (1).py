import os
import re
import html
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import datetime
from tkinter.scrolledtext import ScrolledText
import webbrowser

# 检查必要的库
try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("警告: 未安装python-docx库，将使用基本的文件检查")

try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

try:
    import win32com.client
    HAS_WIN32 = True
except ImportError:
    HAS_WIN32 = False


class DarkMarkChecker:
    def __init__(self, company_names=""):
        self.results = {
            "签章": {"passed": True, "issues": []},
            "版面": {"passed": True, "issues": []},
            "排版": {"passed": True, "issues": []},
            "图表": {"passed": True, "issues": []},
            "内容": {"passed": True, "issues": []}
        }
        self.company_names = company_names.split(",") if company_names else []
        self.company_names = [name.strip() for name in self.company_names if name.strip()]
        self.sensitive_positions = []  # 存储敏感词位置信息
        
        # A4纸张尺寸标准（单位：厘米）
        self.A4_WIDTH_CM = 21.0
        self.A4_HEIGHT_CM = 29.7
        # 允许的误差范围（厘米）
        self.TOLERANCE_CM = 0.1
    
    def check_document(self, filepath):
        """主检查函数"""
        try:
            if not HAS_DOCX:
                messagebox.showerror("错误", "未安装python-docx库，无法检查Word文档")
                return None
            
            doc = Document(filepath)
            
            # 执行各项检查
            self._check_signatures(doc)
            self._check_page_layout(doc)
            self._check_formatting(doc)
            self._check_charts_images(doc)
            self._check_content(doc, filepath)
            
            return self.results
            
        except Exception as e:
            messagebox.showerror("错误", f"打开文档时出错: {str(e)}")
            return None
    
    def _check_signatures(self, doc):
        """检查签章"""
        # 检查是否包含图片（可能的签章）
        if len(doc.inline_shapes) > 0:
            self.results["签章"]["passed"] = False
            self.results["签章"]["issues"].append(f"文档包含 {len(doc.inline_shapes)} 个内联形状/图片，可能包含签章")
    
    def _check_page_layout(self, doc):
        """检查页面设置"""
        for section_idx, section in enumerate(doc.sections):
            # 优化纸张大小检查
            width_cm = self._convert_twips_to_cm(section.page_width)
            height_cm = self._convert_twips_to_cm(section.page_height)
            
            # 检查是否为A4尺寸（允许一定误差）
            width_ok = abs(width_cm - self.A4_WIDTH_CM) <= self.TOLERANCE_CM
            height_ok = abs(height_cm - self.A4_HEIGHT_CM) <= self.TOLERANCE_CM
            
            if not (width_ok and height_ok):
                self.results["版面"]["passed"] = False
                issue = f"纸张大小不是A4 (21×29.7厘米)，当前尺寸: {width_cm:.2f}×{height_cm:.2f}厘米"
                if section_idx > 0:
                    issue += f" (第{section_idx+1}节)"
                self.results["版面"]["issues"].append(issue)
            
            # 检查页边距（2.5厘米）
            margins = [
                ("左", section.left_margin),
                ("右", section.right_margin),
                ("上", section.top_margin),
                ("下", section.bottom_margin)
            ]
            
            for name, margin in margins:
                margin_cm = self._convert_twips_to_cm(margin)
                if abs(margin_cm - 2.5) > 0.05:  # 允许0.05厘米误差
                    self.results["版面"]["passed"] = False
                    issue = f"{name}边距不是2.5厘米，当前为{margin_cm:.2f}厘米"
                    if section_idx > 0:
                        issue += f" (第{section_idx+1}节)"
                    self.results["版面"]["issues"].append(issue)
            
            # 检查页眉页脚
            try:
                # 检查是否有页眉
                header = section.header
                if header and header.paragraphs:
                    has_header_content = any(p.text.strip() for p in header.paragraphs)
                    if has_header_content:
                        self.results["版面"]["passed"] = False
                        issue = "存在页眉内容"
                        if section_idx > 0:
                            issue += f" (第{section_idx+1}节)"
                        self.results["版面"]["issues"].append(issue)
                
                # 检查是否有页脚
                footer = section.footer
                if footer and footer.paragraphs:
                    has_footer_content = any(p.text.strip() for p in footer.paragraphs)
                    if has_footer_content:
                        self.results["版面"]["passed"] = False
                        issue = "存在页脚内容"
                        if section_idx > 0:
                            issue += f" (第{section_idx+1}节)"
                        self.results["版面"]["issues"].append(issue)
            except:
                pass  # 如果无法检查页眉页脚，则跳过
    
    def _convert_twips_to_cm(self, twips):
        """将缇(twips)转换为厘米(cm)"""
        # 1缇 = 1/1440英寸，1英寸 = 2.54厘米
        try:
            # 确保输入是数值类型
            if hasattr(twips, 'cm'):  # 如果已经是Cm对象
                return twips.cm
            elif hasattr(twips, 'pt'):  # 如果是Pt对象
                inches = twips.pt / 72.0  # 1pt = 1/72英寸
                cm = inches * 2.54
                return cm
            else:  # 假设是缇(twips)单位
                inches = float(twips) / 1440.0
                cm = inches * 2.54
                return cm
        except Exception as e:
            print(f"转换单位时出错: {e}, 值: {twips}, 类型: {type(twips)}")
            return 0.0
    
    def _check_formatting(self, doc):
        """检查排版格式"""
        paragraphs_checked = 0
        max_paragraphs_to_check = min(200, len(doc.paragraphs))
        
        for i, paragraph in enumerate(doc.paragraphs[:max_paragraphs_to_check]):
            # 跳过空段落
            if not paragraph.text.strip():
                continue
            
            paragraphs_checked += 1
            
            # 检查字号和字体
            has_runs = False
            for run in paragraph.runs:
                if run.text.strip():  # 只检查有内容的run
                    has_runs = True
                    
                    # 检查字号
                    if run.font.size:
                        font_size_pt = run.font.size.pt if hasattr(run.font.size, 'pt') else run.font.size
                        if abs(font_size_pt - 14) > 0.1:  # 四号字=14磅，允许0.1磅误差
                            self.results["排版"]["passed"] = False
                            issue = f"段落{i+1}: 字号不是四号(14磅)，当前为{font_size_pt:.1f}磅"
                            if issue not in self.results["排版"]["issues"]:
                                self.results["排版"]["issues"].append(issue)
                    
                    # 检查字体
                    if run.font.name:
                        font_name = run.font.name.lower()
                        if '宋体' not in font_name and 'simsun' not in font_name:
                            self.results["排版"]["passed"] = False
                            issue = f"段落{i+1}: 字体不是宋体，当前为{run.font.name}"
                            if issue not in self.results["排版"]["issues"]:
                                self.results["排版"]["issues"].append(issue)
            
            # 如果没有有效的runs但有文本，可能是格式问题
            if not has_runs and paragraph.text.strip():
                self.results["排版"]["issues"].append(f"段落{i+1}: 无法检查字体和字号格式")
            
            # 检查段间距
            if paragraph.paragraph_format.space_before:
                space_before_pt = paragraph.paragraph_format.space_before.pt
                if space_before_pt > 0.1:  # 允许微小误差
                    self.results["排版"]["passed"] = False
                    issue = f"段落{i+1}: 段前间距不为0，当前为{space_before_pt:.1f}磅"
                    if issue not in self.results["排版"]["issues"]:
                        self.results["排版"]["issues"].append(issue)
            
            if paragraph.paragraph_format.space_after:
                space_after_pt = paragraph.paragraph_format.space_after.pt
                if space_after_pt > 0.1:  # 允许微小误差
                    self.results["排版"]["passed"] = False
                    issue = f"段落{i+1}: 段后间距不为0，当前为{space_after_pt:.1f}磅"
                    if issue not in self.results["排版"]["issues"]:
                        self.results["排版"]["issues"].append(issue)
            
            # 检查行间距
            if hasattr(paragraph.paragraph_format, 'line_spacing'):
                line_spacing = paragraph.paragraph_format.line_spacing
                if line_spacing:
                    # 行间距可能是倍数、固定值或None
                    if hasattr(line_spacing, 'pt'):
                        line_spacing_pt = line_spacing.pt
                        if abs(line_spacing_pt - 28) > 0.1:  # 固定值28磅，允许0.1磅误差
                            self.results["排版"]["passed"] = False
                            issue = f"段落{i+1}: 行间距不是固定值28磅，当前为{line_spacing_pt:.1f}磅"
                            if issue not in self.results["排版"]["issues"]:
                                self.results["排版"]["issues"].append(issue)
                    elif isinstance(line_spacing, (int, float)):
                        if abs(line_spacing - 28) > 0.1:  # 假设是磅值
                            self.results["排版"]["passed"] = False
                            issue = f"段落{i+1}: 行间距不是固定值28磅，当前为{line_spacing:.1f}"
                            if issue not in self.results["排版"]["issues"]:
                                self.results["排版"]["issues"].append(issue)
            
            # 只检查前100个非空段落
            if paragraphs_checked >= 100:
                break
        
        # 如果文档段落太少，添加提示
        if paragraphs_checked == 0:
            self.results["排版"]["issues"].append("文档中没有可检查的文本段落")
    
    def _check_charts_images(self, doc):
        """检查图表"""
        # 检查图片
        for i, shape in enumerate(doc.inline_shapes):
            self.results["图表"]["issues"].append(f"发现图片/图表 {i+1}: 请人工检查是否为电脑绘制")
        
        # 检查表格
        for i, table in enumerate(doc.tables):
            self.results["图表"]["issues"].append(f"发现表格 {i+1}: 已检查表格内文字格式")
            
            # 检查表格内文字格式（仿宋五号）
            for row_num, row in enumerate(table.rows):
                for cell_num, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.text.strip():
                                # 检查字号（五号=10.5磅）
                                if run.font.size:
                                    font_size_pt = run.font.size.pt if hasattr(run.font.size, 'pt') else run.font.size
                                    if abs(font_size_pt - 10.5) > 0.1:  # 允许0.1磅误差
                                        self.results["图表"]["passed"] = False
                                        issue = f"表格{i+1}行{row_num+1}列{cell_num+1}: 文字不是五号字(10.5磅)，当前为{font_size_pt:.1f}磅"
                                        if issue not in self.results["图表"]["issues"]:
                                            self.results["图表"]["issues"].append(issue)
                                
                                # 检查字体（仿宋）
                                if run.font.name:
                                    font_name = run.font.name.lower()
                                    if '仿宋' not in font_name and 'fangsong' not in font_name:
                                        self.results["图表"]["passed"] = False
                                        issue = f"表格{i+1}行{row_num+1}列{cell_num+1}: 字体不是仿宋，当前为{run.font.name}"
                                        if issue not in self.results["图表"]["issues"]:
                                            self.results["图表"]["issues"].append(issue)
    
    def _check_content(self, doc, filepath):
        """检查内容中的敏感信息"""
        # 基础敏感词模式 - 移除了"投标人"和"技术"
        sensitive_patterns = [
            # 公司名称相关
            (r'[\u4e00-\u9fa5]{2,30}(?:股份有限公司|有限责任公司|集团公司|有限公司|股份公司|集团)', '完整公司名称'),
            
            # 个人信息
            (r'身份证号[:：]?\s*[\dXx]{15,18}', '身份证号'),
            (r'统一社会信用代码[:：]?\s*[A-Z0-9]{18}', '统一社会信用代码'),
            (r'组织机构代码[:：]?\s*[A-Z0-9]{9}', '组织机构代码'),
            (r'手机[:：]?\s*1[3-9]\d{9}', '手机号'),
            (r'电话[:：]?\s*(?:\d{3,4}-)?\d{7,8}', '电话号码'),
            (r'邮箱[:：]?\s*[\w\.\-]+@[\w\.\-]+\.[a-zA-Z]{2,}', '邮箱地址'),
            
            # 项目业绩相关 - 需要更具体的匹配
            (r'项目名称[:：].*[\u4e00-\u9fa5a-zA-Z0-9]{10,}', '详细项目名称'),  # 长项目名称可能包含识别信息
            (r'业绩证明[:：].*[\u4e00-\u9fa5a-zA-Z0-9]{10,}', '业绩证明内容'),
            (r'合同编号[:：].*[\u4e00-\u9fa5a-zA-Z0-9]{6,}', '合同编号'),
            (r'业主单位[:：].*[\u4e00-\u9fa5a-zA-Z0-9]{4,}', '业主单位'),
            
            # 特定关键词组合（可能泄露身份）
            (r'我(?:公司|单位|司).{0,10}(?:承建|施工|建设|设计)', '第一人称公司描述'),
            (r'本(?:公司|单位|司).{0,10}(?:承建|施工|建设|设计)', '本公司描述'),
            
            # 人名检测（中文名通常2-4个字）
            (r'姓名[:：]\s*[\u4e00-\u9fa5]{2,4}', '姓名'),
            (r'负责人[:：]\s*[\u4e00-\u9fa5]{2,4}', '负责人姓名'),
            (r'联系人[:：]\s*[\u4e00-\u9fa5]{2,4}', '联系人姓名'),
            
            # 地址信息（可能包含公司地址）
            (r'地址[:：].*[\u4e00-\u9fa5a-zA-Z0-9]{10,}', '详细地址'),
            (r'注册地址[:：].*[\u4e00-\u9fa5a-zA-Z0-9]{10,}', '注册地址'),
        ]
        
        # 添加用户输入的公司名称（精确匹配）
        for company in self.company_names:
            if company and len(company) >= 2:  # 至少2个字符
                # 转义特殊字符并确保是完整匹配
                escaped_company = re.escape(company)
                # 匹配整个词，前后可以有标点或空格
                pattern = r'(?:^|\s|[^\w\u4e00-\u9fa5])' + escaped_company + r'(?:$|\s|[^\w\u4e00-\u9fa5])'
                sensitive_patterns.append((pattern, f'指定公司: {company}'))
        
        # 清空之前的敏感词位置记录
        self.sensitive_positions = []
        
        # 检查所有段落
        all_text = ""
        for para_num, paragraph in enumerate(doc.paragraphs, 1):
            para_text = paragraph.text
            if not para_text.strip():
                continue
                
            all_text += para_text + "\n"
            
            # 在当前段落中搜索敏感词
            for pattern, pattern_name in sensitive_patterns:
                try:
                    matches = list(re.finditer(pattern, para_text, re.IGNORECASE))
                    for match in matches:
                        start_pos = match.start()
                        end_pos = match.end()
                        sensitive_text = match.group()
                        
                        # 检查是否真的是敏感词（排除常见非敏感词）
                        if self._is_really_sensitive(sensitive_text, pattern_name):
                            # 记录位置信息
                            self.sensitive_positions.append({
                                'para_num': para_num,
                                'start_pos': start_pos,
                                'end_pos': end_pos,
                                'sensitive_text': sensitive_text,
                                'pattern_name': pattern_name,
                                'context': para_text[max(0, start_pos-20):min(len(para_text), end_pos+20)]
                            })
                            
                            # 添加到问题列表
                            self.results["内容"]["passed"] = False
                            issue_text = f"段落{para_num}: 发现'{pattern_name}' - '{sensitive_text[:30]}...'"
                            if issue_text not in self.results["内容"]["issues"]:
                                self.results["内容"]["issues"].append(issue_text)
                except re.error as e:
                    print(f"正则表达式错误: {e}, 模式: {pattern}")
        
        # 检查表格内容
        for table_num, table in enumerate(doc.tables, 1):
            for row_num, row in enumerate(table.rows, 1):
                for cell_num, cell in enumerate(row.cells, 1):
                    cell_text = cell.text
                    if not cell_text.strip():
                        continue
                        
                    for para_num, paragraph in enumerate(cell.paragraphs, 1):
                        para_text = paragraph.text
                        if not para_text.strip():
                            continue
                            
                        all_text += para_text + "\n"
                        
                        # 在表格段落中搜索敏感词
                        for pattern, pattern_name in sensitive_patterns:
                            try:
                                matches = list(re.finditer(pattern, para_text, re.IGNORECASE))
                                for match in matches:
                                    start_pos = match.start()
                                    end_pos = match.end()
                                    sensitive_text = match.group()
                                    
                                    # 检查是否真的是敏感词
                                    if self._is_really_sensitive(sensitive_text, pattern_name):
                                        # 记录位置信息
                                        self.sensitive_positions.append({
                                            'para_num': f"表格{table_num}行{row_num}列{cell_num}段{para_num}",
                                            'start_pos': start_pos,
                                            'end_pos': end_pos,
                                            'sensitive_text': sensitive_text,
                                            'pattern_name': pattern_name,
                                            'context': para_text[max(0, start_pos-20):min(len(para_text), end_pos+20)]
                                        })
                                        
                                        # 添加到问题列表
                                        self.results["内容"]["passed"] = False
                                        issue_text = f"表格{table_num}行{row_num}列{cell_num}: 发现'{pattern_name}' - '{sensitive_text[:30]}...'"
                                        if issue_text not in self.results["内容"]["issues"]:
                                            self.results["内容"]["issues"].append(issue_text)
                            except re.error as e:
                                print(f"正则表达式错误: {e}, 模式: {pattern}")
        
        # 检查文件名是否包含敏感信息
        filename = os.path.basename(filepath)
        company_keywords = ['公司', '有限', '责任', '股份', '集团', '投标', '标书', '技术标']
        for keyword in company_keywords:
            if keyword in filename:
                self.results["内容"]["passed"] = False
                issue = f"文件名可能包含敏感信息: {filename} (包含'{keyword}')"
                if issue not in self.results["内容"]["issues"]:
                    self.results["内容"]["issues"].append(issue)
        
        # 检查用户输入的公司名称是否在文件名中
        for company in self.company_names:
            if company and company in filename:
                self.results["内容"]["passed"] = False
                issue = f"文件名包含公司名称: {company}"
                if issue not in self.results["内容"]["issues"]:
                    self.results["内容"]["issues"].append(issue)
                break
        
        # 如果没有发现任何敏感词，添加通过提示
        if not self.sensitive_positions and self.results["内容"]["passed"]:
            self.results["内容"]["issues"].append("✅ 未发现敏感信息")
    
    def _is_really_sensitive(self, text, pattern_name):
        """判断是否真的是敏感词（排除误判）"""
        # 通用非敏感词列表
        non_sensitive_words = [
            '技术', '投标', '文件', '文档', '格式', '要求', '检查',
            '系统', '软件', '程序', '代码', '功能', '模块', '界面',
            '用户', '操作', '设置', '配置', '参数', '选项', '按钮'
        ]
        
        # 排除模式
        exclude_patterns = [
            r'^[大小长短高低]$',  # 单个形容词
            r'^\d+$',  # 纯数字
            r'^[a-zA-Z]{1,3}$',  # 短英文
            r'^[\u4e00-\u9fa5]{1}$',  # 单个汉字
        ]
        
        text_clean = text.strip()
        
        # 只排除完全等于通用词的内容，避免漏检“某某技术有限公司”等真实敏感信息
        if text_clean in non_sensitive_words:
            return False
        
        # 检查排除模式
        for pattern in exclude_patterns:
            if re.match(pattern, text):
                return False
        
        # 针对特定模式进行额外检查
        if pattern_name == '公司类型词':
            # 如果只是单独的公司类型词，没有具体名称，不算敏感
            if len(text) <= 4 and text in ['公司', '有限公司', '集团公司']:
                return False
        
        return True
    
    def get_sensitive_positions(self):
        """获取敏感词位置信息"""
        return self.sensitive_positions
    
    def generate_highlighted_text(self, doc):
        """生成高亮显示敏感词的文本（HTML格式）"""
        if not self.sensitive_positions:
            return ""
        
        html_content = """
        <html>
        <head>
        <meta charset="UTF-8">
        <style>
            body { font-family: Arial, sans-serif; line-height: 1.6; margin: 20px; }
            .highlight { background-color: #ffcccc; font-weight: bold; padding: 2px; border-radius: 3px; }
            .para-header { background-color: #f0f0f0; padding: 10px; margin: 15px 0; font-weight: bold; border-left: 4px solid #2196F3; }
            .sensitive-info { color: #d32f2f; font-weight: bold; margin: 10px 0 5px 0; }
            .context { border: 1px solid #ddd; padding: 10px; margin: 10px 0; background-color: #f9f9f9; border-radius: 5px; }
            .position { color: #666; font-size: 0.9em; }
            .summary { background-color: #e8f5e8; padding: 15px; margin: 20px 0; border-radius: 5px; }
            h2 { color: #333; border-bottom: 2px solid #4CAF50; padding-bottom: 10px; }
            .warning { color: #ff9800; }
            table { border-collapse: collapse; width: 100%; margin: 10px 0; }
            th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
            th { background-color: #f2f2f2; }
        </style>
        </head>
        <body>
        <h2>技术标暗标敏感词检测报告</h2>
        """
        
        # 添加摘要信息
        company_summary = html.escape(', '.join(self.company_names) if self.company_names else '无')
        html_content += f"""
        <div class="summary">
            <strong>检查摘要：</strong><br>
            检测到 <span style="color:red; font-weight:bold;">{len(self.sensitive_positions)}</span> 处可能敏感词<br>
            检查时间：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}<br>
            排查公司：{company_summary}
        </div>
        """
        
        # 按类型统计
        type_stats = {}
        for pos in self.sensitive_positions:
            type_name = pos['pattern_name']
            type_stats[type_name] = type_stats.get(type_name, 0) + 1
        
        if type_stats:
            html_content += '<div class="para-header">📊 敏感词类型统计</div>'
            html_content += '<table>'
            html_content += '<tr><th>敏感词类型</th><th>数量</th></tr>'
            for type_name, count in sorted(type_stats.items()):
                html_content += f'<tr><td>{type_name}</td><td>{count}</td></tr>'
            html_content += '</table>'
        
        # 按段落分组
        positions_by_para = {}
        for pos in self.sensitive_positions:
            para_key = pos['para_num']
            if para_key not in positions_by_para:
                positions_by_para[para_key] = []
            positions_by_para[para_key].append(pos)
        
        # 生成每个段落的高亮文本
        for para_num, positions in positions_by_para.items():
            html_content += f'<div class="para-header">📄 位置: {html.escape(str(para_num))}</div>'
            
            # 获取段落文本
            para_text = ""
            if isinstance(para_num, int):
                if 0 <= para_num-1 < len(doc.paragraphs):
                    para_text = doc.paragraphs[para_num-1].text
            
            # 高亮显示敏感词
            if para_text:
                sorted_positions = sorted(positions, key=lambda x: x['start_pos'])
                highlighted_parts = []
                cursor = 0
                for pos in sorted_positions:
                    start = pos.get('start_pos', 0)
                    end = pos.get('end_pos', 0)
                    if start < cursor or end > len(para_text) or start >= end:
                        continue
                    highlighted_parts.append(html.escape(para_text[cursor:start]))
                    highlighted_parts.append(f'<span class="highlight">{html.escape(para_text[start:end])}</span>')
                    cursor = end
                highlighted_parts.append(html.escape(para_text[cursor:]))
                html_content += f'<div class="context">{"".join(highlighted_parts)}</div>'
            
            # 显示检测到的敏感词详情
            html_content += '<div class="sensitive-info">🔍 检测到的敏感词：</div>'
            for pos in positions:
                pattern_name = html.escape(str(pos["pattern_name"]))
                sensitive_text = html.escape(str(pos["sensitive_text"]))
                html_content += f"""
                <div>
                    • <strong>{pattern_name}</strong>: "{sensitive_text}" 
                    <span class="position">(位置: {pos["start_pos"]}-{pos["end_pos"]})</span>
                </div>
                """
            
            html_content += '<hr>'
        
        # 添加建议
        html_content += """
        <div class="para-header">💡 修改建议</div>
        <div class="context">
            <strong>注意：以下为可能敏感信息，请人工复核</strong><br><br>
            <strong>建议操作：</strong><br>
            1. 检查并删除或替换所有公司名称、单位名称<br>
            2. 移除所有个人姓名、联系方式<br>
            3. 删除具体项目业绩中的可识别信息<br>
            4. 移除所有合同编号、证件号码等唯一标识<br>
            5. 检查图表中是否包含敏感信息<br>
            6. 确保文件名不包含公司、投标等关键词<br>
            7. 重新检查格式要求（字体、字号、页边距等）<br><br>
            
            <strong>特别注意：</strong><br>
            • "技术"、"投标"等通用词不被视为敏感词<br>
            • 需要删除的是具体可识别身份的信息<br>
            • 所有修改需保持文档的完整性和可读性
        </div>
        """
        
        html_content += '</body></html>'
        return html_content


class DarkMarkCheckerGUI:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("技术标暗标格式核查软件 v2.3")
        self.root.geometry("1180x780")
        self.root.minsize(980, 680)
        self.root.resizable(True, True)
        self.colors = {
            "bg": "#f4f7fb",
            "card": "#ffffff",
            "primary": "#2563eb",
            "primary_dark": "#1d4ed8",
            "success": "#16a34a",
            "warning": "#f59e0b",
            "danger": "#dc2626",
            "muted": "#64748b",
            "border": "#dbe3ef"
        }
        self.root.configure(bg=self.colors["bg"])
        
        # 设置图标
        try:
            self.root.iconbitmap(default='')  # 可以添加图标文件路径
        except:
            pass
        
        # 检查必要的库
        self.check_dependencies()
        
        # 设置UI
        self.setup_ui()
        
        # 检查器实例
        self.checker = None
        self.current_doc = None
        self.current_file = None
        
    def setup_styles(self):
        style = ttk.Style()
        try:
            style.theme_use("clam")
        except tk.TclError:
            pass
        style.configure("TNotebook", background=self.colors["bg"], borderwidth=0)
        style.configure("TNotebook.Tab", padding=(16, 8), font=("微软雅黑", 9))
        style.map("TNotebook.Tab", background=[("selected", self.colors["card"])], foreground=[("selected", self.colors["primary_dark"])])
        style.configure("Treeview", rowheight=28, font=("微软雅黑", 9), background=self.colors["card"], fieldbackground=self.colors["card"], bordercolor=self.colors["border"])
        style.configure("Treeview.Heading", font=("微软雅黑", 9, "bold"), background="#e8eef8", foreground="#1e293b")
        style.configure("Vertical.TScrollbar", background=self.colors["border"], troughcolor=self.colors["bg"])
        style.configure("Horizontal.TProgressbar", troughcolor="#e2e8f0", background=self.colors["primary"], bordercolor=self.colors["border"], lightcolor=self.colors["primary"], darkcolor=self.colors["primary"])
    
    def check_dependencies(self):
        """检查依赖库"""
        missing_libs = []
        if not HAS_DOCX:
            missing_libs.append("python-docx (用于处理Word文档)")
        
        if missing_libs:
            message = "以下Python库未安装，部分功能可能受限：\n\n"
            message += "\n".join(missing_libs)
            message += "\n\n是否继续运行？"
            
            response = messagebox.askyesno("警告", message)
            if not response:
                self.root.destroy()
                import sys
                sys.exit(0)
    
    def setup_ui(self):
        self.setup_styles()
        # 创建主框架
        main_frame = tk.Frame(self.root, bg=self.colors["bg"])
        main_frame.pack(fill=tk.BOTH, expand=True, padx=14, pady=14)
        
        # 标题
        title_frame = tk.Frame(main_frame, bg=self.colors["primary"], bd=0, highlightthickness=0)
        title_frame.pack(fill=tk.X, pady=(0, 12), ipady=12)
        
        title_label = tk.Label(title_frame, text="📋 技术标暗标格式核查系统", 
                              font=("微软雅黑", 18, "bold"), bg=self.colors["primary"], fg="white")
        title_label.pack(side=tk.LEFT, padx=18)
        
        version_label = tk.Label(title_frame, text="v2.3", 
                                font=("微软雅黑", 10, "bold"), bg=self.colors["primary"], fg="#dbeafe")
        version_label.pack(side=tk.RIGHT, padx=18)
        
        # 创建左右两个主要区域
        paned_window = tk.PanedWindow(main_frame, orient=tk.HORIZONTAL, sashwidth=5)
        paned_window.pack(fill=tk.BOTH, expand=True, pady=10)
        
        left_frame = tk.Frame(paned_window)
        right_frame = tk.Frame(paned_window)
        
        paned_window.add(left_frame, width=500)
        paned_window.add(right_frame, width=600)
        
        # ========== 左侧区域：控制和检查结果 ==========
        # 控制面板
        control_frame = tk.LabelFrame(left_frame, text="控制面板", font=("微软雅黑", 10, "bold"), bg=self.colors["card"], fg="#1e293b", bd=1, relief=tk.GROOVE, padx=8, pady=8)
        control_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # 公司名称输入区域
        company_frame = tk.Frame(control_frame)
        company_frame.pack(fill=tk.X, padx=10, pady=5)
        
        tk.Label(company_frame, text="排查公司名称（关键词）:", 
                font=("微软雅黑", 9)).pack(anchor=tk.W)
        
        self.company_var = tk.StringVar()
        company_entry = tk.Entry(company_frame, textvariable=self.company_var, 
                                font=("微软雅黑", 10), relief=tk.SOLID, bd=1, highlightthickness=1, highlightcolor=self.colors["primary"], highlightbackground=self.colors["border"])
        company_entry.pack(fill=tk.X, pady=(0, 5))
        company_entry.insert(0, "有限公司,集团公司,股份公司")
        
        tk.Label(company_frame, text="多个关键词用逗号分隔，如：公司A,公司B", 
                font=("微软雅黑", 8), fg=self.colors["muted"], bg=self.colors["card"]).pack(anchor=tk.W)
        
        # 文件操作区域
        file_btn_frame = tk.Frame(control_frame, bg=self.colors["card"])
        file_btn_frame.pack(fill=tk.X, padx=10, pady=5)
        
        self.select_btn = tk.Button(file_btn_frame, text="📂 选择Word文件", 
                                   command=self.select_file,
                                   font=("微软雅黑", 10),
                                   bg=self.colors["success"], fg="white",
                                   activebackground="#15803d", activeforeground="white", relief=tk.FLAT, cursor="hand2",
                                   padx=15, pady=8)
        self.select_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        self.check_btn = tk.Button(file_btn_frame, text="🔍 开始检查", 
                                  command=self.check_file,
                                  font=("微软雅黑", 10),
                                  bg="#2196F3", fg="white",
                                  padx=15, pady=8,
                                  state=tk.DISABLED)
        self.check_btn.pack(side=tk.LEFT)
        
        # 文件路径显示
        self.file_path_var = tk.StringVar()
        path_frame = tk.Frame(control_frame, bg=self.colors["card"])
        path_frame.pack(fill=tk.X, padx=10, pady=(0, 10))
        
        self.path_label = tk.Label(path_frame, textvariable=self.file_path_var,
                                  fg=self.colors["primary_dark"], bg=self.colors["card"], font=("微软雅黑", 9), anchor=tk.W, 
                                  wraplength=440, justify=tk.LEFT)
        self.path_label.pack(fill=tk.X)
        
        # 检查结果区域
        result_frame = tk.LabelFrame(left_frame, text="检查结果", font=("微软雅黑", 11, "bold"), bg=self.colors["card"], fg="#1e293b", bd=1, relief=tk.GROOVE, padx=8, pady=8)
        result_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # 创建Treeview显示结果
        self.tree = ttk.Treeview(result_frame, columns=("项目", "状态", "问题"), 
                                show="headings", height=12)
        
        # 设置列
        self.tree.heading("项目", text="检查项目", anchor=tk.W)
        self.tree.heading("状态", text="检查状态", anchor=tk.W)
        self.tree.heading("问题", text="发现问题", anchor=tk.W)
        
        self.tree.column("项目", width=80, anchor=tk.W)
        self.tree.column("状态", width=80, anchor=tk.W)
        self.tree.column("问题", width=320, anchor=tk.W)
        
        # 添加滚动条
        scrollbar = ttk.Scrollbar(result_frame, orient=tk.VERTICAL, 
                                 command=self.tree.yview)
        self.tree.configure(yscrollcommand=scrollbar.set)
        
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(5, 0), pady=5)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y, pady=5)
        
        # 绑定Treeview选择事件
        self.tree.bind('<<TreeviewSelect>>', self.on_tree_select)
        
        # 底部操作区域
        bottom_frame = tk.Frame(left_frame, bg=self.colors["bg"])
        bottom_frame.pack(fill=tk.X, padx=5, pady=10)
        
        # 汇总结果标签
        self.summary_label = tk.Label(bottom_frame, text="", 
                                     font=("微软雅黑", 10, "bold"), bg=self.colors["bg"], fg="#1e293b")
        self.summary_label.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        # 操作按钮
        self.export_btn = tk.Button(bottom_frame, text="📄 导出报告", 
                                   command=self.export_report,
                                   font=("微软雅黑", 9, "bold"),
                                   bg=self.colors["warning"], fg="white",
                                   activebackground="#d97706", activeforeground="white", relief=tk.FLAT, cursor="hand2",
                                   padx=15, pady=6,
                                   state=tk.DISABLED)
        self.export_btn.pack(side=tk.RIGHT, padx=(5, 0))
        
        # ========== 右侧区域：详情显示 ==========
        # 详情标签页
        notebook = ttk.Notebook(right_frame)
        notebook.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Tab 1: 敏感词高亮显示
        highlight_tab = tk.Frame(notebook, bg=self.colors["card"])
        notebook.add(highlight_tab, text="敏感词高亮")
        
        # 创建带滚动条的文本显示区域
        self.highlight_text = ScrolledText(highlight_tab, wrap=tk.WORD, 
                                          font=("微软雅黑", 9),
                                          height=25)
        self.highlight_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # 设置文本标签样式
        self.highlight_text.tag_configure("highlight", background="#ffcccc", foreground="red")
        self.highlight_text.tag_configure("header", font=("微软雅黑", 10, "bold"))
        self.highlight_text.tag_configure("sensitive", foreground="red", font=("微软雅黑", 9, "bold"))
        
        # Tab 2: 检查要求
        requirement_tab = tk.Frame(notebook)
        notebook.add(requirement_tab, text="检查要求")
        
        req_text = """📌 技术标暗标格式要求：

1. ✅ 签章要求
   • 不得对"暗标"文件进行电子签章
   • 不得上传带有签章的扫描页

2. 📄 版面要求
   • 采用A4纸张大小 (21×29.7厘米)
   • 页边距均为2.5厘米
   • 不得出现页眉、页脚、页码

3. ✍️ 排版要求
   • 全文均为白底黑字
   • 字体：宋体
   • 字号：四号 (14磅)
   • 段间距：0
   • 行间距：固定值28磅

4. 📊 图表要求
   • 应电脑绘制，不得手绘
   • 白底黑字
   • 文字采用仿宋五号 (10.5磅)

5. 🔒 内容要求（特别注意）
   • 不得出现投标人名称和其他可识别投标人身份的字符、徽标
   • 不得出现企业名称、过往项目业绩内容
   • 不得出现投标人独有的标准名称或编号
   • 不得出现其他具有标识性作用的符号、图案等
   
   ❗ 注意以下词不被视为敏感词：
   • "技术" - 技术标文件中的正常词汇
   • "投标" - 投标文件中的正常词汇
   • 其他通用技术术语

⚠️ 重要提示：
对投标文件未按上述"暗标"要求制作的，作无效投标处理"""
        
        req_label = tk.Label(requirement_tab, text=req_text, justify=tk.LEFT,
                            font=("微软雅黑", 9), anchor=tk.NW)
        req_label.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Tab 3: 使用说明
        help_tab = tk.Frame(notebook, bg=self.colors["card"])
        notebook.add(help_tab, text="使用说明")
        
        help_text = """📖 使用说明：

1. 🔍 检查前准备
   • 输入需要排查的公司名称关键词，多个用逗号分隔
   • 示例：有限公司,集团公司,股份公司,公司A

2. 📂 选择文件
   • 点击"选择Word文件"按钮
   • 选择要检查的Word文档(.docx)

3. 🔍 开始检查
   • 点击"开始检查"按钮
   • 程序将自动检查所有格式要求

4. 📊 查看结果
   • 左侧显示检查结果汇总
   • "敏感词高亮"标签显示敏感词位置
   • 红色背景表示检测到的可能敏感词

5. 📄 导出报告
   • 点击"导出报告"保存详细检查结果
   • 可以查看HTML格式的高亮报告

💡 敏感词检测说明：
• 程序会检测：公司名称、个人姓名、联系方式、项目业绩等
• 不会检测："技术"、"投标"等通用词汇
• 检测结果需要人工复核确认

⚙️ 高级功能：
• 支持多个章节的文档检查
• 支持表格内容检查
• 支持精确的格式检查（字体、字号、间距）
• 支持自定义公司名称排查

⚠️ 注意事项：
• 此工具为辅助工具，检查结果需要人工复核
• 部分格式问题可能需要人工确认
• 建议在最终提交前多次检查"""
        
        help_label = tk.Label(help_tab, text=help_text, justify=tk.LEFT,
                             font=("微软雅黑", 9), anchor=tk.NW, bg=self.colors["card"], fg="#1e293b")
        help_label.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 操作按钮区域
        action_frame = tk.Frame(right_frame, bg=self.colors["bg"])
        action_frame.pack(fill=tk.X, padx=5, pady=5)
        
        self.show_html_btn = tk.Button(action_frame, text="🌐 查看HTML报告", 
                                      command=self.show_html_report,
                                      font=("微软雅黑", 9),
                                      bg="#9C27B0", fg="white",
                                      padx=15, pady=6,
                                      state=tk.DISABLED)
        self.show_html_btn.pack(side=tk.LEFT, padx=(0, 5))
        
        self.clear_btn = tk.Button(action_frame, text="🗑️ 清空显示", 
                                  command=self.clear_highlight,
                                  font=("微软雅黑", 9, "bold"),
                                  bg=self.colors["muted"], fg="white",
                                  activebackground="#475569", activeforeground="white", relief=tk.FLAT, cursor="hand2",
                                  padx=15, pady=6)
        self.clear_btn.pack(side=tk.LEFT)
        
        self.quick_check_btn = tk.Button(action_frame, text="⚡ 快速检查", 
                                        command=self.quick_check,
                                        font=("微软雅黑", 9, "bold"),
                                        bg=self.colors["success"], fg="white",
                                        activebackground="#15803d", activeforeground="white", relief=tk.FLAT, cursor="hand2",
                                        padx=15, pady=6)
        self.quick_check_btn.pack(side=tk.RIGHT)
        
        # 状态栏
        status_frame = tk.Frame(self.root, bg="#e2e8f0", relief=tk.FLAT, bd=0, height=28)
        status_frame.pack(side=tk.BOTTOM, fill=tk.X)
        status_frame.pack_propagate(False)
        
        self.status_label = tk.Label(status_frame, text="就绪", 
                                    font=("微软雅黑", 8), anchor=tk.W, bg="#e2e8f0", fg="#334155")
        self.status_label.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        
        # 进度条
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(status_frame, variable=self.progress_var, 
                                           mode='determinate', length=100)
        self.progress_bar.pack(side=tk.RIGHT, padx=5)
        
        # 绑定窗口关闭事件
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)
        
    def select_file(self):
        """选择Word文件"""
        self.update_status("正在选择文件...")
        
        filepath = filedialog.askopenfilename(
            title="选择Word文档",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        
        if filepath:
            if not filepath.lower().endswith('.docx'):
                messagebox.showwarning("文件格式不支持", "当前版本仅支持 .docx 文件，请先将 .doc 文件另存为 .docx 后再检查。")
                self.update_status("文件格式不支持")
                return
            self.file_path_var.set(filepath)
            self.check_btn.config(state=tk.NORMAL)
            self.current_file = filepath
            
            # 清空之前的结果
            self.clear_results()
            
            self.update_status(f"已选择文件: {os.path.basename(filepath)}")
            self.progress_var.set(0)
        else:
            self.update_status("已取消选择")
    
    def clear_results(self):
        """清空检查结果"""
        for item in self.tree.get_children():
            self.tree.delete(item)
        self.summary_label.config(text="")
        self.export_btn.config(state=tk.DISABLED)
        self.show_html_btn.config(state=tk.DISABLED)
        self.clear_highlight()
        self.progress_var.set(0)
    
    def check_file(self):
        """执行检查"""
        if not self.current_file:
            return
            
        # 清空之前的检查结果
        self.clear_results()
        
        self.update_status("正在检查文档...")
        self.progress_var.set(10)
        self.root.update()
        
        try:
            # 创建检查器，传入公司名称
            company_names = self.company_var.get().strip()
            self.checker = DarkMarkChecker(company_names)
            
            # 执行检查
            self.progress_var.set(30)
            self.root.update()
            
            results = self.checker.check_document(self.current_file)
            
            if results:
                # 保存文档对象用于高亮显示
                try:
                    self.current_doc = Document(self.current_file)
                except Exception as e:
                    self.current_doc = None
                    print(f"无法保存文档对象: {e}")
                
                self.progress_var.set(70)
                self.root.update()
                
                total_passed = 0
                total_failed = 0
                
                for category, data in results.items():
                    status = "✅ 通过" if data["passed"] else "❌ 失败"
                    
                    # 插入主项
                    item_id = self.tree.insert("", "end", values=(category, status, ""))
                    
                    # 插入问题详情
                    if data["passed"]:
                        total_passed += 1
                    else:
                        total_failed += 1
                    
                    if data["issues"]:
                        for issue in data["issues"][:8]:  # 最多显示8个问题
                            self.tree.insert(item_id, "end", values=("", "", f"• {issue}"))
                
                self.progress_var.set(90)
                self.root.update()
                
                # 显示汇总结果
                if total_failed == 0:
                    self.summary_label.config(text="🎉 所有检查项目通过！", fg="green")
                else:
                    self.summary_label.config(
                        text=f"⚠️ {total_failed}个项目未通过，{total_passed}个项目通过", 
                        fg="red"
                    )
                
                self.export_btn.config(state=tk.NORMAL)
                self.show_html_btn.config(state=tk.NORMAL)
                
                # 显示敏感词高亮信息
                self.show_sensitive_highlights()
                
                self.progress_var.set(100)
                self.update_status(f"检查完成 - 共发现{len(self.checker.sensitive_positions)}处可能敏感词")
                
        except Exception as e:
            messagebox.showerror("检查错误", f"检查过程中出错: {str(e)}")
            self.update_status("检查出错")
            self.progress_var.set(0)
    
    def quick_check(self):
        """快速检查（仅检查内容和基本格式）"""
        if not self.current_file:
            self.select_file()
            return
            
        self.clear_results()
        self.update_status("正在快速检查...")
        self.progress_var.set(20)
        
        try:
            company_names = self.company_var.get().strip()
            self.checker = DarkMarkChecker(company_names)
            
            if HAS_DOCX:
                doc = Document(self.current_file)
                
                # 快速检查纸张大小
                for section in doc.sections:
                    width_cm = self.checker._convert_twips_to_cm(section.page_width)
                    height_cm = self.checker._convert_twips_to_cm(section.page_height)
                    
                    if not (abs(width_cm - 21.0) <= 0.2 and abs(height_cm - 29.7) <= 0.2):
                        self.tree.insert("", "end", values=("版面", "❌ 失败", "• 纸张大小可能不是A4"))
                        break
                
                # 检查内容
                self.checker._check_content(doc, self.current_file)
                self.current_doc = doc
                
                # 显示结果
                if self.checker.sensitive_positions:
                    self.summary_label.config(text=f"⚠️ 发现{len(self.checker.sensitive_positions)}处可能敏感词", fg="orange")
                    self.show_sensitive_highlights()
                else:
                    self.summary_label.config(text="✅ 未发现可能敏感词", fg="green")
                
                self.export_btn.config(state=tk.NORMAL)
                self.show_html_btn.config(state=tk.NORMAL)
                self.progress_var.set(100)
                self.update_status("快速检查完成")
                
        except Exception as e:
            messagebox.showerror("快速检查错误", f"快速检查出错: {str(e)}")
            self.update_status("快速检查出错")
            self.progress_var.set(0)
    
    def show_sensitive_highlights(self):
        """在右侧显示敏感词高亮信息"""
        if not self.checker or not self.current_doc:
            return
        
        # 清空文本区域
        self.highlight_text.delete(1.0, tk.END)
        
        # 获取敏感词位置信息
        positions = self.checker.get_sensitive_positions()
        if not positions:
            self.highlight_text.insert(tk.END, "✅ 未发现可能敏感词\n", "header")
            return
        
        self.highlight_text.insert(tk.END, f"🔍 共发现 {len(positions)} 处可能敏感词（需要人工复核）\n\n", "header")
        
        # 按类型统计
        type_stats = {}
        for pos in positions:
            type_name = pos['pattern_name']
            type_stats[type_name] = type_stats.get(type_name, 0) + 1
        
        if type_stats:
            self.highlight_text.insert(tk.END, "📊 敏感词类型统计：\n", "header")
            for type_name, count in sorted(type_stats.items()):
                self.highlight_text.insert(tk.END, f"  {type_name}: {count}处\n")
            self.highlight_text.insert(tk.END, "\n")
        
        # 按段落分组
        positions_by_para = {}
        for pos in positions:
            para_key = pos['para_num']
            if para_key not in positions_by_para:
                positions_by_para[para_key] = []
            positions_by_para[para_key].append(pos)
        
        # 显示每个段落的敏感词
        for para_num, positions in positions_by_para.items():
            self.highlight_text.insert(tk.END, f"\n📄 位置: {para_num}\n", "header")
            
            # 获取段落文本
            para_text = ""
            if isinstance(para_num, int):
                if 0 <= para_num-1 < len(self.current_doc.paragraphs):
                    para_text = self.current_doc.paragraphs[para_num-1].text
            
            if para_text:
                # 显示段落文本，高亮敏感词
                start_idx = self.highlight_text.index(tk.END)
                self.highlight_text.insert(tk.END, para_text + "\n")
                end_idx = self.highlight_text.index(tk.END)
                
                # 高亮每个敏感词
                for pos in positions:
                    if 'start_pos' in pos and 'end_pos' in pos:
                        # 计算在文本控件中的位置
                        try:
                            line_start = f"{start_idx}+{pos['start_pos']}c"
                            line_end = f"{start_idx}+{pos['end_pos']}c"
                            self.highlight_text.tag_add("highlight", line_start, line_end)
                        except:
                            pass
            
            # 显示检测详情
            self.highlight_text.insert(tk.END, "\n🔍 检测到的可能敏感词:\n", "sensitive")
            for pos in positions[:5]:  # 只显示前5个
                detail = f"  • {pos['pattern_name']}: '{pos['sensitive_text']}' "
                detail += f"(位置: {pos['start_pos']}-{pos['end_pos']})\n"
                self.highlight_text.insert(tk.END, detail)
            
            if len(positions) > 5:
                self.highlight_text.insert(tk.END, f"  ... 还有 {len(positions)-5} 个可能敏感词\n")
            
            self.highlight_text.insert(tk.END, "-" * 60 + "\n")
        
        # 添加说明
        self.highlight_text.insert(tk.END, "\n💡 说明：\n", "header")
        self.highlight_text.insert(tk.END, "• '技术'、'投标'等通用词不会被标记为敏感词\n")
        self.highlight_text.insert(tk.END, "• 以上检测结果需要人工复核确认\n")
        self.highlight_text.insert(tk.END, "• 红色高亮部分为可能敏感信息\n")
    
    def on_tree_select(self, event):
        """Treeview选择事件"""
        selection = self.tree.selection()
        if selection:
            item = self.tree.item(selection[0])
            values = item['values']
            if values and len(values) > 0 and values[0]:
                # 可以在这里添加更多处理逻辑
                pass
    
    def show_html_report(self):
        """显示HTML格式的高亮报告"""
        if not self.checker or not self.current_doc:
            return
        
        # 生成HTML内容
        html_content = self.checker.generate_highlighted_text(self.current_doc)
        
        if not html_content:
            messagebox.showinfo("提示", "未发现可能敏感词，无需生成高亮报告")
            return
        
        # 保存为临时HTML文件
        temp_file = os.path.join(os.path.expanduser("~"), "Desktop", f"暗标检查报告_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.html")
        try:
            with open(temp_file, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            # 在默认浏览器中打开
            webbrowser.open(f"file://{temp_file}")
            self.update_status(f"HTML报告已生成: {os.path.basename(temp_file)}")
            
        except Exception as e:
            messagebox.showerror("错误", f"生成HTML报告时出错: {str(e)}")
    
    def clear_highlight(self):
        """清空高亮显示区域"""
        self.highlight_text.delete(1.0, tk.END)
    
    def export_report(self):
        """导出检查报告"""
        if not self.current_file:
            return
            
        save_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("文本文件", "*.txt"), ("HTML文件", "*.html"), ("所有文件", "*.*")],
            initialfile=f"暗标检查报告_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}"
        )
        
        if save_path:
            try:
                if save_path.endswith('.html'):
                    # 导出HTML报告
                    if self.checker and self.current_doc:
                        html_content = self.checker.generate_highlighted_text(self.current_doc)
                        with open(save_path, 'w', encoding='utf-8') as f:
                            f.write(html_content)
                        messagebox.showinfo("成功", f"HTML报告已保存到:\n{save_path}")
                else:
                    # 导出文本报告
                    with open(save_path, 'w', encoding='utf-8') as f:
                        f.write("=" * 70 + "\n")
                        f.write("技术标暗标格式检查报告\n")
                        f.write("=" * 70 + "\n")
                        f.write(f"检查文件: {self.current_file}\n")
                        f.write(f"检查时间: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                        f.write(f"排查公司关键词: {self.company_var.get() or '无'}\n")
                        f.write("=" * 70 + "\n\n")
                        
                        results = self.checker.results if self.checker else {}
                        for category, data in results.items():
                            status = "通过" if data["passed"] else "失败"
                            f.write(f"{category}: {status}\n")
                            
                            if data["issues"]:
                                for issue in data["issues"]:
                                    f.write(f"  • {issue}\n")
                            f.write("\n")
                        
                        # 添加敏感词详细位置信息
                        if self.checker and self.checker.sensitive_positions:
                            f.write("\n" + "=" * 70 + "\n")
                            f.write("可能敏感词详细位置信息（需要人工复核）\n")
                            f.write("=" * 70 + "\n")
                            
                            # 按类型统计
                            type_stats = {}
                            for pos in self.checker.sensitive_positions:
                                type_name = pos['pattern_name']
                                type_stats[type_name] = type_stats.get(type_name, 0) + 1
                            
                            if type_stats:
                                f.write("\n敏感词类型统计:\n")
                                for type_name, count in sorted(type_stats.items()):
                                    f.write(f"  {type_name}: {count}处\n")
                                f.write("\n")
                            
                            # 按段落分组
                            positions_by_para = {}
                            for pos in self.checker.sensitive_positions:
                                para_key = pos['para_num']
                                if para_key not in positions_by_para:
                                    positions_by_para[para_key] = []
                                positions_by_para[para_key].append(pos)
                            
                            for para_num, positions in positions_by_para.items():
                                f.write(f"\n位置: {para_num}\n")
                                for pos in positions:
                                    f.write(f"  • {pos['pattern_name']}: '{pos['sensitive_text']}' ")
                                    f.write(f"(位置: {pos['start_pos']}-{pos['end_pos']})\n")
                                    f.write(f"    上下文: ...{pos['context']}...\n")
                            
                            f.write("\n" + "=" * 70 + "\n")
                            f.write("重要说明：\n")
                            f.write("1. '技术'、'投标'等通用词不会被标记为敏感词\n")
                            f.write("2. 以上检测结果需要人工复核确认\n")
                            f.write("3. 重点检查公司名称、个人姓名、联系方式等\n")
                    
                    messagebox.showinfo("成功", f"检查报告已保存到:\n{save_path}")
                
                self.update_status(f"报告已导出: {os.path.basename(save_path)}")
                
            except Exception as e:
                messagebox.showerror("错误", f"保存报告时出错: {str(e)}")
    
    def update_status(self, message):
        """更新状态栏"""
        self.status_label.config(text=f"状态: {message}")
        self.root.update()
    
    def on_closing(self):
        """窗口关闭事件"""
        if messagebox.askokcancel("退出", "确定要退出程序吗？"):
            self.root.destroy()
    
    def run(self):
        """运行GUI"""
        # 居中显示窗口
        self.root.update_idletasks()
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        window_width = 1180
        window_height = 780
        x = (screen_width // 2) - (window_width // 2)
        y = (screen_height // 2) - (window_height // 2)
        self.root.geometry(f'{window_width}x{window_height}+{x}+{y}')
        
        self.root.mainloop()


def show_installation_guide():
    """显示安装指南"""
    print("=" * 70)
    print("技术标暗标核查软件 - 安装指南")
    print("=" * 70)
    print("\n请先安装以下Python库：")
    print("1. python-docx (处理Word文档) - 必需")
    print("2. Pillow (图像处理) - 可选")
    print("3. pywin32 (Windows COM接口) - 可选")
    print("\n安装命令：")
    print("pip install python-docx")
    print("\n或完整安装：")
    print("pip install python-docx Pillow pywin32")
    print("\n如果遇到安装问题，可以尝试指定版本：")
    print("pip install python-docx==0.8.11")
    print("=" * 70)


if __name__ == "__main__":
    import sys
    
    # 检查是否安装了必要的库
    if not HAS_DOCX:
        show_installation_guide()
        response = input("\n缺少必要库，是否继续运行？(y/n): ")
        if response.lower() != 'y':
            sys.exit(1)
    
    try:
        # 运行GUI应用
        app = DarkMarkCheckerGUI()
        app.run()
    except Exception as e:
        messagebox.showerror("错误", f"程序启动失败: {str(e)}")
        print(f"错误信息: {e}")
        import traceback
        traceback.print_exc()