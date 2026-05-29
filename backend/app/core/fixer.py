"""暗标格式修复模块 v1.0.4 - 增强样式和段落字体修复"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, Any
from pathlib import Path

EMU_PER_CM = 360000
EMU_PER_PT = 12700


def _same_emu(actual, expected, tolerance=1000):
    if actual is None:
        return False
    return abs(int(actual) - int(expected)) <= tolerance


class DarkBidFixer:
    """暗标格式修复器"""

    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.req = config.get("requirements", {})

    def fix(self, input_path: str, output_path: str) -> Dict[str, Any]:
        """修复文档格式"""
        doc = Document(input_path)
        changes = []

        changes.extend(self._fix_page_settings(doc))
        changes.extend(self._fix_header_footer(doc))
        changes.extend(self._fix_text_format(doc))
        changes.extend(self._fix_table_format(doc))
        changes.extend(self._fix_styles(doc))

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

        return {
            "success": True,
            "output_path": output_path,
            "total_changes": len(changes),
            "changes": changes,
        }

    def _fix_page_settings(self, doc: Document) -> list:
        """修复页面设置"""
        changes = []
        page_req = self.req.get("page", {})
        margins = page_req.get("margins", {})

        width_mm = page_req.get("width_mm", 210)
        height_mm = page_req.get("height_mm", 297)
        page_size = page_req.get("size", "A4")

        changed = False
        expected_width = Cm(width_mm / 10)
        expected_height = Cm(height_mm / 10)
        expected_margins = {
            "top_margin": Cm(margins.get("top_cm", 2.5)),
            "bottom_margin": Cm(margins.get("bottom_cm", 2.5)),
            "left_margin": Cm(margins.get("left_cm", 2.5)),
            "right_margin": Cm(margins.get("right_cm", 2.5)),
        }

        for section in doc.sections:
            if not _same_emu(section.page_width, expected_width):
                section.page_width = expected_width
                changed = True
            if not _same_emu(section.page_height, expected_height):
                section.page_height = expected_height
                changed = True
            for attr, expected in expected_margins.items():
                if not _same_emu(getattr(section, attr), expected):
                    setattr(section, attr, expected)
                    changed = True

        if changed:
            changes.append(f"页面设置: 纸张大小 {page_size}, 页边距已按配置修复")
        return changes

    def _fix_header_footer(self, doc: Document) -> list:
        """修复页眉页脚"""
        changes = []
        page_req = self.req.get("page", {})

        if page_req.get("noHeader", True):
            changed = False
            for section in doc.sections:
                changed = self._clear_container(section.header) or changed
            if changed:
                changes.append("删除页眉")

        if page_req.get("noFooter", True):
            changed = False
            for section in doc.sections:
                changed = self._clear_container(section.footer) or changed
            if changed:
                changes.append("删除页脚")

        return changes

    def _clear_container(self, container):
        changed = False
        if container.is_linked_to_previous:
            has_content = self._container_has_content(container)
            if not has_content:
                return False
            changed = True
        else:
            has_content = self._container_has_content(container)
            changed = has_content
        container.is_linked_to_previous = False
        for child in list(container._element):
            if child.tag.endswith('}p') or child.tag.endswith('}tbl'):
                if child.tag.endswith('}tbl') or ''.join(child.itertext()).strip() or "PAGE" in child.xml or "NUMPAGES" in child.xml:
                    changed = True
                container._element.remove(child)
        container.add_paragraph()
        return changed

    @staticmethod
    def _container_has_content(container):
        for para in container.paragraphs:
            if para.text.strip() or "PAGE" in para._element.xml or "NUMPAGES" in para._element.xml:
                return True
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        return True
        return False

    def _fix_text_format(self, doc: Document) -> list:
        """修复正文格式"""
        changes = []
        text_req = self.req.get("text", {})

        expected_font = text_req.get("font", "宋体")
        expected_size = text_req.get("fontSize_pt", 14)
        expected_color = text_req.get("fontColor", "000000")
        expected_spacing = text_req.get("paragraphSpacing_pt", 0)
        expected_line_spacing = text_req.get("lineSpacing", {}).get("value_pt", 28)

        changed = self._fix_paragraphs(
            doc.paragraphs,
            expected_font,
            expected_size,
            expected_color,
            expected_spacing,
            expected_line_spacing,
        )

        if changed["font"]:
            changes.append(f"字体: 修改为 {expected_font}")
        if changed["size"]:
            changes.append(f"字号: 修改为 {expected_size}pt")
        if changed["color"]:
            changes.append(f"字体颜色: 修改为 #{expected_color}")
        if changed["spacing"]:
            changes.append(f"段落间距: 段前段后 {expected_spacing}pt, 行距 {expected_line_spacing}pt")
        if changed["format"]:
            changes.append("格式: 已清除加粗/倾斜/下划线/删除线")

        return changes

    def _fix_table_format(self, doc: Document) -> list:
        """修复表格格式"""
        changes = []
        table_req = self.req.get("table", {})
        text_req = self.req.get("text", {})

        expected_font = table_req.get("font", "仿宋")
        expected_size = table_req.get("fontSize_pt", 10.5)
        expected_color = table_req.get("fontColor", "000000")
        expected_spacing = text_req.get("paragraphSpacing_pt", 0)
        expected_line_spacing = text_req.get("lineSpacing", {}).get("value_pt", 28)

        all_paragraphs = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_paragraphs.extend(cell.paragraphs)

        changed = self._fix_paragraphs(
            all_paragraphs,
            expected_font,
            expected_size,
            expected_color,
            expected_spacing,
            expected_line_spacing,
        )

        if changed["font"]:
            changes.append(f"表格字体: 修改为 {expected_font}")
        if changed["size"]:
            changes.append(f"表格字号: 修改为 {expected_size}pt")
        if changed["color"]:
            changes.append(f"表格字体颜色: 修改为 #{expected_color}")
        if changed["spacing"]:
            changes.append(f"表格段落间距: 段前段后 {expected_spacing}pt, 行距 {expected_line_spacing}pt")
        if changed["format"]:
            changes.append("表格格式: 已清除加粗/倾斜/下划线/删除线")

        return changes

    def _fix_styles(self, doc: Document) -> list:
        expected_font = self.req.get("text", {}).get("font", "宋体")
        changed = False
        for style in doc.styles:
            if not hasattr(style, "_element") or not hasattr(style, "font"):
                continue
            rPr = style._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            if self._fix_rfonts(rFonts, expected_font):
                changed = True
            if style.font.name != expected_font:
                style.font.name = expected_font
                changed = True
        defaults = doc.styles.element.find(qn('w:docDefaults'))
        if defaults is not None:
            rPr_default = defaults.find(qn('w:rPrDefault'))
            rPr = rPr_default.find(qn('w:rPr')) if rPr_default is not None else None
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = rPr._add_rFonts()
                if self._fix_rfonts(rFonts, expected_font):
                    changed = True
        return [f"样式字体: 修改为 {expected_font}"] if changed else []

    @staticmethod
    def _fix_rfonts(rFonts, expected_font):
        changed = False
        for attr in ['ascii', 'hAnsi', 'eastAsia', 'cs']:
            key = qn(f'w:{attr}')
            if rFonts.get(key) != expected_font:
                rFonts.set(key, expected_font)
                changed = True
        for attr in ['asciiTheme', 'eastAsiaTheme', 'hAnsiTheme', 'cstheme']:
            key = qn(f'w:{attr}')
            if key in rFonts.attrib:
                del rFonts.attrib[key]
                changed = True
        return changed

    def _fix_paragraphs(self, paragraphs, expected_font, expected_size, expected_color, expected_spacing, expected_line_spacing):
        changed = {"font": False, "size": False, "color": False, "spacing": False, "format": False}
        color_rgb = RGBColor.from_string(expected_color)

        for para in paragraphs:
            pf = para.paragraph_format
            pPr = para._element.get_or_add_pPr()
            p_rPr = pPr.find(qn('w:rPr'))
            if p_rPr is None:
                p_rPr = OxmlElement('w:rPr')
                pPr.insert(0, p_rPr)
            p_rFonts = p_rPr.find(qn('w:rFonts'))
            if p_rFonts is None:
                p_rFonts = OxmlElement('w:rFonts')
                p_rPr.insert(0, p_rFonts)
            if self._fix_rfonts(p_rFonts, expected_font):
                changed["font"] = True
            if not _same_emu(pf.space_before, Pt(expected_spacing)):
                pf.space_before = Pt(expected_spacing)
                changed["spacing"] = True
            if not _same_emu(pf.space_after, Pt(expected_spacing)):
                pf.space_after = Pt(expected_spacing)
                changed["spacing"] = True
            if pf.line_spacing_rule != WD_LINE_SPACING.EXACTLY or not _same_emu(pf.line_spacing, Pt(expected_line_spacing)):
                pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                pf.line_spacing = Pt(expected_line_spacing)
                changed["spacing"] = True

            for run in para.runs:
                # 检查是否需要修复字体（检查 eastAsia 和 font.name）
                need_fix_font = False
                rPr = run._element.rPr
                if rPr is None or rPr.rFonts is None:
                    need_fix_font = True
                elif any(rPr.rFonts.get(qn(f'w:{attr}')) != expected_font for attr in ['ascii', 'hAnsi', 'eastAsia', 'cs']):
                    need_fix_font = True
                if not need_fix_font and run.font.name != expected_font:
                    need_fix_font = True

                if need_fix_font:
                    run.font.name = expected_font
                    rFonts = run._element.get_or_add_rPr().get_or_add_rFonts()
                    self._fix_rfonts(rFonts, expected_font)
                    changed["font"] = True
                if not _same_emu(run.font.size, Pt(expected_size)):
                    run.font.size = Pt(expected_size)
                    changed["size"] = True
                if run.font.color is None or run.font.color.rgb != color_rgb:
                    run.font.color.rgb = color_rgb
                    changed["color"] = True
                # 清除格式：加粗/倾斜/下划线/删除线
                if run.font.bold:
                    run.font.bold = False
                    changed["format"] = True
                if run.font.italic:
                    run.font.italic = False
                    changed["format"] = True
                if run.font.underline is not None:
                    run.font.underline = None
                    changed["format"] = True
                if run.font.strike:
                    run.font.strike = False
                    changed["format"] = True

        return changed
