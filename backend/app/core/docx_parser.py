"""Word 文档解析模块"""
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from typing import Dict, List, Any, Optional


class DocxParser:
    """解析 Word 文档结构和格式"""

    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = Document(file_path)
        self._paragraphs_info = None
        self._tables_info = None

    def get_page_settings(self) -> Dict[str, Any]:
        """获取首页页面设置信息"""
        return self.get_sections_info()[0]

    def get_sections_info(self) -> List[Dict[str, Any]]:
        """获取所有节的页面设置信息"""
        sections = []
        for index, section in enumerate(self.doc.sections):
            sections.append({
                "index": index,
                "page_width_mm": self._emu_to_mm(section.page_width),
                "page_height_mm": self._emu_to_mm(section.page_height),
                "margin_top_cm": self._emu_to_cm(section.top_margin),
                "margin_bottom_cm": self._emu_to_cm(section.bottom_margin),
                "margin_left_cm": self._emu_to_cm(section.left_margin),
                "margin_right_cm": self._emu_to_cm(section.right_margin),
                "has_header": self._has_header(section),
                "has_footer": self._has_footer(section),
                "has_page_number": self._section_has_page_number(section),
            })
        return sections

    def get_paragraphs_info(self) -> List[Dict[str, Any]]:
        """获取所有段落的格式信息"""
        if self._paragraphs_info is not None:
            return self._paragraphs_info

        self._paragraphs_info = []
        for i, para in enumerate(self.doc.paragraphs):
            if not para.text.strip() and not para.runs:
                continue

            para_info = {
                "index": i,
                "text": para.text[:100] if para.text else "",
                "style_name": para.style.name if para.style else "Normal",
                "runs": [],
                "paragraph_format": self._get_paragraph_format(para),
            }

            for run in para.runs:
                para_info["runs"].append(self._get_run_info(run, para))

            self._paragraphs_info.append(para_info)

        return self._paragraphs_info

    def get_tables_info(self) -> List[Dict[str, Any]]:
        """获取所有表格的格式信息"""
        if self._tables_info is not None:
            return self._tables_info

        self._tables_info = []
        for i, table in enumerate(self.doc.tables):
            table_info = {
                "index": i,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "cells": [],
            }

            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cell_info = {
                        "row": row_idx,
                        "col": col_idx,
                        "text": cell.text[:50] if cell.text else "",
                        "paragraphs": [],
                    }

                    for para in cell.paragraphs:
                        para_info = {
                            "text": para.text[:30] if para.text else "",
                            "paragraph_format": self._get_paragraph_format(para),
                            "runs": [],
                        }
                        for run in para.runs:
                            para_info["runs"].append(self._get_run_info(run, para))
                        cell_info["paragraphs"].append(para_info)

                    table_info["cells"].append(cell_info)

            self._tables_info.append(table_info)

        return self._tables_info

    def get_all_text(self) -> str:
        """获取文档所有文本内容"""
        texts = []
        for para in self.doc.paragraphs:
            if para.text.strip():
                texts.append(para.text)
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        texts.append(cell.text)
        return "\n".join(texts)

    def has_images(self) -> bool:
        """检查文档是否包含图片"""
        for rel in self.doc.part.rels.values():
            if "image" in rel.reltype:
                return True
        return False

    def has_embedded_objects(self) -> bool:
        """检查文档是否包含嵌入对象"""
        for rel in self.doc.part.rels.values():
            reltype = rel.reltype.lower()
            if "oleobject" in reltype or "package" in reltype:
                return True
        return False

    def _get_paragraph_format(self, para) -> Dict[str, Any]:
        """获取段落格式详情"""
        pf = para.paragraph_format
        result = {
            "alignment": str(pf.alignment) if pf.alignment else None,
            "space_before_pt": self._emu_to_pt(pf.space_before) if pf.space_before else 0,
            "space_after_pt": self._emu_to_pt(pf.space_after) if pf.space_after else 0,
            "line_spacing": None,
            "line_spacing_rule": None,
        }

        if pf.line_spacing is not None:
            if pf.line_spacing_rule == WD_LINE_SPACING.EXACTLY:
                result["line_spacing"] = self._emu_to_pt(pf.line_spacing)
                result["line_spacing_rule"] = "exact"
            elif pf.line_spacing_rule == WD_LINE_SPACING.AT_LEAST:
                result["line_spacing"] = self._emu_to_pt(pf.line_spacing)
                result["line_spacing_rule"] = "at_least"
            elif pf.line_spacing_rule == WD_LINE_SPACING.MULTIPLE:
                result["line_spacing"] = pf.line_spacing / 12700
                result["line_spacing_rule"] = "multiple"

        return result

    def _get_run_info(self, run, para) -> Dict[str, Any]:
        return {
            "text": run.text[:50] if run.text else "",
            "font_name": self._effective_font_name(run, para),
            "font_size_pt": self._effective_font_size(run, para),
            "font_color": self._effective_font_color(run, para),
            "bold": run.font.bold,
            "italic": run.font.italic,
        }

    def _effective_font_name(self, run, para) -> Optional[str]:
        if run.font.name:
            return run.font.name
        if run.style and run.style.font.name:
            return run.style.font.name
        if para.style and para.style.font.name:
            return para.style.font.name
        return None

    def _effective_font_size(self, run, para) -> Optional[float]:
        if run.font.size:
            return self._emu_to_pt(run.font.size)
        if run.style and run.style.font.size:
            return self._emu_to_pt(run.style.font.size)
        if para.style and para.style.font.size:
            return self._emu_to_pt(para.style.font.size)
        return None

    def _effective_font_color(self, run, para) -> Optional[str]:
        color = self._get_color_hex(run.font.color)
        if color:
            return color
        if run.style:
            color = self._get_color_hex(run.style.font.color)
            if color:
                return color
        if para.style:
            color = self._get_color_hex(para.style.font.color)
            if color:
                return color
        return None

    @staticmethod
    def _emu_to_mm(emu: Optional[int]) -> Optional[float]:
        """EMU 转毫米"""
        if emu is None:
            return None
        return emu / 914400 * 25.4

    @staticmethod
    def _emu_to_cm(emu: Optional[int]) -> Optional[float]:
        """EMU 转厘米"""
        if emu is None:
            return None
        return emu / 360000

    @staticmethod
    def _emu_to_pt(emu: Optional[int]) -> Optional[float]:
        """EMU 转磅"""
        if emu is None:
            return None
        return emu / 12700

    @staticmethod
    def _get_color_hex(color) -> Optional[str]:
        """获取颜色的十六进制值"""
        if color is None or color.rgb is None:
            return None
        return str(color.rgb)

    def _has_header(self, section) -> bool:
        """检查是否有页眉"""
        try:
            return self._container_has_content(section.header)
        except Exception:
            return False

    def _has_footer(self, section) -> bool:
        """检查是否有页脚"""
        try:
            return self._container_has_content(section.footer)
        except Exception:
            return False

    def _container_has_content(self, container) -> bool:
        for para in container.paragraphs:
            if para.text.strip() or self._paragraph_has_page_number(para):
                return True
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        return True
                    for para in cell.paragraphs:
                        if self._paragraph_has_page_number(para):
                            return True
        return False

    def _section_has_page_number(self, section) -> bool:
        containers = [section.header, section.footer]
        return any(self._container_has_page_number(container) for container in containers)

    def _container_has_page_number(self, container) -> bool:
        for para in container.paragraphs:
            if self._paragraph_has_page_number(para):
                return True
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if self._paragraph_has_page_number(para):
                            return True
        return False

    @staticmethod
    def _paragraph_has_page_number(para) -> bool:
        xml = para._element.xml
        return "PAGE" in xml or "NUMPAGES" in xml
